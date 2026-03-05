import os
import re
from pathlib import Path
from typing import List, Literal, Optional

import openpyxl
import xlwings as xw
from docx import Document
import fitz

OUTPUT_DIR: Path = Path(os.getcwd()) / "output_docs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def _get_file_path(filename: str, default_ext: str = ".docx") -> Path:
    safe_filename: str = os.path.basename(filename)
    if not Path(safe_filename).suffix:
        safe_filename += default_ext
    return OUTPUT_DIR / safe_filename

def list_local_documents() -> str:
    if not OUTPUT_DIR.exists():
        return "Directory does not exist."
    files = [f.name for f in OUTPUT_DIR.iterdir() if f.is_file() and f.suffix in ['.docx', '.xlsx', '.txt', '.pdf']]
    return f"Documents: {', '.join(files)}" if files else "No documents found."

def init_document(filename: str) -> str:
    filepath = _get_file_path(filename, ".docx")
    doc = Document()
    doc.save(filepath)
    return f"Document initialized: {filepath.name}"

def _add_markdown_paragraph(doc: Document, text: str, style: Optional[str] = None) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph(style=style)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def append_to_document(
    filename: str, 
    content: str, 
    element_type: Literal["paragraph", "heading", "list_item"] = "paragraph",
    level: int = 1
) -> str:
    filepath = _get_file_path(filename, ".docx")
    
    try:
        doc = Document(filepath) if filepath.exists() else Document()
    except Exception:
        return "Error: Cannot open file. Is it open in Word?"
        
    if element_type == "heading":
        clean_content = content.replace('**', '')
        doc.add_heading(clean_content, level=level)
    elif element_type == "list_item":
        _add_markdown_paragraph(doc, content, style='List Bullet')
    else:
        _add_markdown_paragraph(doc, content)
        
    try:
        doc.save(filepath)
        return f"Element appended to {filepath.name}."
    except PermissionError:
        return "Error: File is locked. Close Microsoft Word."

def replace_in_document(filename: str, old_text: str, new_text: str) -> str:
    filepath = _get_file_path(filename, ".docx")
    if not filepath.exists():
        return "Error: File not found."
        
    try:
        doc = Document(filepath)
        replaced = False
        for p in doc.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
                replaced = True
                
        if replaced:
            doc.save(filepath)
            return f"Modification successful in {filepath.name}."
        return f"Error: Text '{old_text}' not found."
    except PermissionError:
        return "Error: File is locked."
    except Exception as e:
        return f"Unexpected error: {e}"

def _read_docx(filepath: Path) -> str:
    try:
        doc = Document(filepath)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(lines) if lines else "Document is empty."
    except Exception as e:
        return f"Error reading docx: {e}"

def _read_txt(filepath: Path) -> str:
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"Error reading txt: {e}"

def list_excel_sheets(filename: str) -> str:
    filepath = _get_file_path(filename, ".xlsx")
    if not filepath.exists():
        return "Error: File not found."
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        sheets = wb.sheetnames
        wb.close()
        return f"Sheets in {filepath.name}: {', '.join(sheets)}"
    except Exception as e:
        return f"Error: {e}"

def read_excel_sheet(filename: str, sheet_name: Optional[str] = None, max_rows: int = 50) -> str:
    filepath = _get_file_path(filename, ".xlsx")
    if not filepath.exists():
        return "Error: File not found."
        
    try:
        # Double chargement : valeurs en cache et formules brutes
        wb_val = openpyxl.load_workbook(filepath, data_only=True)
        wb_form = openpyxl.load_workbook(filepath, data_only=False)
        
        if sheet_name and sheet_name in wb_val.sheetnames:
            ws_val = wb_val[sheet_name]
        else:
            ws_val = wb_val.active
            
        ws_form = wb_form[ws_val.title]
        
        data = []
        for r_idx, (row_val, row_form) in enumerate(zip(ws_val.iter_rows(), ws_form.iter_rows())):
            if r_idx >= max_rows:
                data.append("... (truncated)")
                break
                
            row_str_parts = []
            is_empty = True
            
            for cell_val, cell_form in zip(row_val, row_form):
                # 1. Valeur calculée disponible (ou texte standard)
                if cell_val.value is not None:
                    row_str_parts.append(str(cell_val.value).strip())
                    is_empty = False
                # 2. Valeur non calculée mais formule présente
                elif cell_form.value is not None and str(cell_form.value).startswith('='):
                    row_str_parts.append(f"[Formule non évaluée: {cell_form.value}]")
                    is_empty = False
                # 3. Cellule vide
                else:
                    row_str_parts.append("")
                    
            if not is_empty:
                data.append(" | ".join(row_str_parts))
                
        wb_val.close()
        wb_form.close()
        
        if not data:
            return f"Sheet '{ws_val.title}' is empty."
            
        return f"Sheet '{ws_val.title}' content:\n" + "\n".join(data)
        
    except Exception as e:
        return f"Error: {e}"
    
def _refresh_excel_cache(filepath: Path) -> None:
    """Ouvre Excel pour calculer et sauvegarder les formules (Mode Debug)."""
    try:
        import xlwings as xw
        import logging
        logger = logging.getLogger(__name__)
        
        abs_path = str(filepath.resolve())
        logger.info(f"🐛 [DEBUG] Tentative d'ouverture via xlwings : {abs_path}")
        
        # Le mode visible=True permet de voir si Excel affiche une popup bloquante
        with xw.App(visible=False) as app:
            app.display_alerts = False
            wb = app.books.open(abs_path)
            
            # Force le calcul au cas où Excel serait configuré en calcul manuel
            app.calculate()
            
            wb.save()
            wb.close()
            
        logger.info("🐛 [DEBUG] xlwings a calculé et sauvegardé le fichier avec succès.")
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"❌ [DEBUG CRITIQUE] Échec de xlwings : {e}")

def read_document_unified(filename: str, sheet_name: Optional[str] = None) -> str:
    safe_filename: str = os.path.basename(filename)
    filepath: Path = OUTPUT_DIR / safe_filename
    
    if not filepath.exists():
        return f"Error: {safe_filename} not found."
        
    ext: str = filepath.suffix.lower()
    if ext == ".docx":
        return _read_docx(filepath)
    elif ext == ".txt":
        return _read_txt(filepath)
    elif ext == ".xlsx":
        return read_excel_sheet(safe_filename, sheet_name)
    
    return f"Error: Unsupported format {ext}."

def write_to_excel(filename: str, sheet_name: str, data: List[List[str]], start_row: int = 1, start_col: int = 1) -> str:
    """
    Crée ou modifie un fichier Excel en y insérant une grille de données (liste de listes).
    """
    filepath = _get_file_path(filename, ".xlsx")
    
    try:
        if filepath.exists():
            wb = openpyxl.load_workbook(filepath)
            # Sélectionne la feuille existante ou en crée une nouvelle
            if sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
            else:
                ws = wb.create_sheet(sheet_name)
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = sheet_name

        # Insertion des données
        for r_idx, row in enumerate(data, start=start_row):
            for c_idx, value in enumerate(row, start=start_col):
                # Tente de convertir en int/float si possible pour qu'Excel reconnaisse les nombres
                try:
                    if '.' in str(value):
                        value = float(value)
                    else:
                        value = int(value)
                except ValueError:
                    pass # Garde en string si ce n'est pas un nombre
                    
                ws.cell(row=r_idx, column=c_idx, value=value)

        wb.save(filepath)
        wb.close()
        _refresh_excel_cache(filepath)
        return f"Données insérées avec succès dans {filepath.name} (Feuille: {sheet_name})."
    
    except PermissionError:
        return "Erreur : Le fichier est ouvert dans Excel. Fermez-le pour le modifier."
    except Exception as e:
        return f"Erreur inattendue lors de l'écriture Excel : {e}"
    
def refresh_excel_file(filename: str) -> str:
    """Ouvre temporairement Excel pour calculer toutes les formules d'un fichier."""
    filepath = _get_file_path(filename, ".xlsx")
    if not filepath.exists():
        return f"Erreur : Le fichier {filepath.name} n'existe pas."
        
    try:
        _refresh_excel_cache(filepath)
        return f"Succès : Le fichier {filepath.name} a été actualisé. Les formules sont maintenant calculées, tu peux utiliser l'outil de lecture pour voir les résultats."
    except Exception as e:
        return f"Erreur lors de l'actualisation : {e}"
    
def read_pdf(filename: str) -> str:
    path = OUTPUT_DIR / filename
    if not path.exists() or not path.is_file():
        return f"Error: File {filename} not found."
        
    try:
        text: list[str] = []
        with fitz.open(path) as doc:
            for page in doc:
                text.append(page.get_text())
        return "\n".join(text).strip()
    except Exception as e:
        return f"Error reading PDF: {e}"